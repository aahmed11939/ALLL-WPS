"""
Word / DOCX design report builder — ALLL WPS Designer.

Produces a stamped engineering memorandum with:
  - Title page with project metadata and stamp block
  - Executive summary
  - Basis of design / system description
  - Hydraulic calculations (Darcy-Weisbach, head budget, system curve)
  - Pump analysis (H-Q / η / P / NPSHr curves, operating points, VFD)
  - Wet well sizing and pump cycling (AWWA M32)
  - Engineering checks summary
  - Surge / water hammer analysis (Joukowsky + MOC)
  - Surge protection device comparison
  - Appendices: full inputs, pipe schedule, K-value reference

Figures are generated with matplotlib and embedded as PNG streams.
"""

from __future__ import annotations

import io
import math
from datetime import date as _date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Lazy matplotlib import — avoids import overhead if unused
# ---------------------------------------------------------------------------

def _plt():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    return plt


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

NAVY   = RGBColor(0x1F, 0x38, 0x64)
TEAL   = RGBColor(0x1A, 0x52, 0x76)
GREY   = RGBColor(0x4A, 0x4A, 0x4A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

FIGURE_WIDTH_IN = 5.8   # inches


# ---------------------------------------------------------------------------
# Unit helpers
# ---------------------------------------------------------------------------

def _us(unit_system: str, val: float | None, kind: str) -> tuple[float | None, str]:
    """Convert SI value → display unit. Returns (converted_val, unit_label)."""
    if val is None:
        return None, "—"
    if unit_system != "US":
        labels = {
            "flow": "m³/h", "head": "m", "pressure": "kPa",
            "velocity": "m/s", "length": "m", "volume": "m³",
            "power": "kW", "temp": "°C",
        }
        return val, labels.get(kind, "—")
    conv = {
        "flow":     (val * 4.40287,          "gpm"),
        "head":     (val * 3.28084,          "ft"),
        "pressure": (val * 0.14504,          "psi"),
        "velocity": (val * 3.28084,          "ft/s"),
        "length":   (val * 3.28084,          "ft"),
        "volume":   (val * 264.172,          "gal"),
        "power":    (val * 1.34102,          "hp"),
        "temp":     (val * 9 / 5 + 32,       "°F"),
    }
    return conv.get(kind, (val, "—"))


def _fmt(val: float | None, dp: int = 2) -> str:
    if val is None:
        return "—"
    return f"{val:,.{dp}f}"


# ---------------------------------------------------------------------------
# Document style helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "AAAAAA") -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = TEAL
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = GREY
        run.font.size = Pt(11)


def _para(doc: Document, text: str, bold: bool = False,
          italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _kv_table(doc: Document, rows: list[tuple[str, str]],
              col_widths: tuple[float, float] = (2.5, 3.5)) -> None:
    """Compact 2-column key-value table."""
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i]
        key_cell = row.cells[0]
        val_cell = row.cells[1]
        key_cell.text = k
        val_cell.text = str(v)
        key_cell.paragraphs[0].runs[0].bold = True
        key_cell.paragraphs[0].runs[0].font.size = Pt(9)
        val_cell.paragraphs[0].runs[0].font.size  = Pt(9)
        _set_cell_bg(key_cell, "D0E4F5")
        _set_cell_borders(key_cell)
        _set_cell_borders(val_cell)
        key_cell.width = Inches(col_widths[0])
        val_cell.width = Inches(col_widths[1])


def _col_table(doc: Document, headers: list[str],
               data_rows: list[list[str]],
               col_widths: list[float] | None = None) -> None:
    """Multi-column data table with styled header row."""
    if not data_rows:
        return
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = h
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        _set_cell_bg(cell, "1F3864")
        _set_cell_borders(cell)
        if col_widths:
            cell.width = Inches(col_widths[j])
    for i, dr in enumerate(data_rows):
        row = tbl.rows[i + 1]
        bg  = "F2F5F8" if i % 2 == 1 else "FFFFFF"
        for j, val in enumerate(dr):
            cell = row.cells[j]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell)


def _embed_figure(doc: Document, fig_bytes: bytes, caption: str = "",
                  width_in: float = FIGURE_WIDTH_IN) -> None:
    """Embed a PNG byte stream as a centred figure."""
    buf = io.BytesIO(fig_bytes)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _hline(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Figure generators
# ---------------------------------------------------------------------------

def _fig_system_pump_curve(draft: dict) -> bytes | None:
    """Generate pump H-Q + system curve overlay PNG."""
    hyd   = draft.get("hydraulicsResult") or {}
    pump  = draft.get("pumpResult")       or {}

    sys_pts  = hyd.get("system_curve", [])  or []
    hq_pts   = pump.get("hq_curve", [])     or []
    ops      = pump.get("operating_points", []) or []
    sp_curves= pump.get("speed_curves", []) or []

    if not sys_pts and not hq_pts:
        return None

    plt = _plt()
    fig, ax = plt.subplots(figsize=(6, 3.8))

    if sys_pts:
        qs = [p.get("Q_m3h", p.get("Q_m3h", 0)) for p in sys_pts]
        hs = [p.get("H_m", 0) for p in sys_pts]
        ax.plot(qs, hs, color="#2E5B8C", linewidth=2, label="System curve")

    if hq_pts:
        qs = [p.get("Q_m3h", 0) for p in hq_pts]
        hs = [p.get("value", 0) for p in hq_pts]
        ax.plot(qs, hs, color="#1A5276", linewidth=2.5, label="Pump H-Q (rated)")

    for sc in sp_curves[:5]:
        spd  = sc.get("speed_pct", 100)
        pts  = sc.get("hq_pts", [])
        if pts:
            qs2 = [p.get("Q_m3h", 0) for p in pts]
            hs2 = [p.get("value", p.get("H_m", 0)) for p in pts]
            ax.plot(qs2, hs2, "--", linewidth=1.2,
                    color="#5D8AA8", label=f"{spd}% speed")

    for op in ops:
        ax.scatter([op.get("Q_m3h", 0)], [op.get("H_m", 0)],
                   color="#E74C3C", zorder=5,
                   label=f"OP ({op.get('n_pumps', 1)} pump)")

    ax.set_xlabel("Flow Q (m³/h)", fontsize=9)
    ax.set_ylabel("Head H (m)",    fontsize=9)
    ax.set_title("Pump H-Q and System Curve",  fontsize=10, fontweight="bold")
    ax.legend(fontsize=8, loc="best")
    ax.grid(True, alpha=0.3)
    ax.set_xlim(left=0)
    ax.set_ylim(bottom=0)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _fig_efficiency_curves(draft: dict) -> bytes | None:
    """Generate efficiency / power / NPSHr trio figure."""
    pump = draft.get("pumpResult") or {}
    eta_pts   = pump.get("eta_curve", [])   or []
    p_pts     = pump.get("p_curve", [])     or []
    npshr_pts = pump.get("npshr_curve", []) or []

    if not eta_pts:
        return None

    plt = _plt()
    fig, axes = plt.subplots(1, 3, figsize=(9, 3.2))
    datasets = [
        (axes[0], eta_pts,   "Efficiency η (%)",  "#27AE60"),
        (axes[1], p_pts,     "Power P (kW)",      "#E67E22"),
        (axes[2], npshr_pts, "NPSHr (m)",         "#8E44AD"),
    ]
    for ax, pts, ylabel, color in datasets:
        if pts:
            qs = [p.get("Q_m3h", 0) for p in pts]
            vs = [p.get("value", 0) for p in pts]
            ax.plot(qs, vs, color=color, linewidth=2)
            ax.fill_between(qs, vs, alpha=0.15, color=color)
        ax.set_xlabel("Q (m³/h)", fontsize=8)
        ax.set_ylabel(ylabel, fontsize=8)
        ax.grid(True, alpha=0.3)
        ax.set_xlim(left=0)
        ax.set_ylim(bottom=0)

    fig.suptitle("Pump Performance Curves", fontsize=10, fontweight="bold")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _fig_surge_envelope(draft: dict) -> bytes | None:
    """Plot MOC pressure envelope (H_max, H_min, pipe elevation) vs distance."""
    moc = draft.get("mocResult") or {}
    env = moc.get("envelope", []) or []

    if not env:
        return None

    plt = _plt()
    fig, ax = plt.subplots(figsize=(6, 3.5))

    xs      = [p.get("x_m", 0)     for p in env]
    h_max   = [p.get("H_max_m", 0) for p in env]
    h_min   = [p.get("H_min_m", 0) for p in env]
    elev    = [p.get("elev_m", 0)  for p in env]

    ax.plot(xs, h_max, color="#C0392B", linewidth=2, label="H max (transient)")
    ax.plot(xs, h_min, color="#2980B9", linewidth=2, label="H min (transient)")
    ax.plot(xs, elev,  "k--",          linewidth=1, label="Pipe elevation")
    ax.fill_between(xs, h_min, h_max, alpha=0.1, color="#5D6D7E")

    ax.set_xlabel("Distance from pump (m)", fontsize=9)
    ax.set_ylabel("Head (m)",               fontsize=9)
    ax.set_title("MOC Surge Pressure Envelope", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _fig_moc_time_history(draft: dict) -> bytes | None:
    """Plot MOC head time histories at key observation nodes."""
    moc  = draft.get("mocResult") or {}
    obs  = moc.get("observations", []) or []

    if not obs:
        return None

    plt = _plt()
    fig, ax = plt.subplots(figsize=(6, 3.5))
    colors  = ["#1F3864", "#C0392B", "#27AE60", "#E67E22", "#8E44AD"]

    for idx, ob in enumerate(obs[:5]):
        hist  = ob.get("history", []) or []
        label = ob.get("label", f"Node {idx}")
        if hist:
            ts = [h.get("t_s", 0)  for h in hist]
            hs = [h.get("H_m", 0)  for h in hist]
            ax.plot(ts, hs, color=colors[idx % len(colors)],
                    linewidth=1.8, label=label)

    ax.set_xlabel("Time (s)",   fontsize=9)
    ax.set_ylabel("Head H (m)", fontsize=9)
    ax.set_title("MOC Head Time Histories", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _fig_protection_comparison(draft: dict) -> bytes | None:
    """Bar chart comparing Hmax/Hmin across baseline + device runs."""
    wi = draft.get("whatIfResult") or {}
    baseline = wi.get("baseline") or {}
    runs     = wi.get("device_runs", []) or []

    all_runs = [baseline] + runs
    labels   = [r.get("label", f"Run {i}") for i, r in enumerate(all_runs)]
    h_max    = [r.get("global_max_H_m", 0) or 0 for r in all_runs]
    h_min    = [r.get("global_min_H_m", 0) or 0 for r in all_runs]

    if not labels:
        return None

    plt = _plt()
    x   = range(len(labels))
    fig, ax = plt.subplots(figsize=(max(5, len(labels) * 1.5), 3.5))

    w = 0.35
    ax.bar([i - w / 2 for i in x], h_max, width=w, color="#C0392B", alpha=0.85, label="H max")
    ax.bar([i + w / 2 for i in x], h_min, width=w, color="#2980B9", alpha=0.85, label="H min")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=8)
    ax.set_ylabel("Head (m)", fontsize=9)
    ax.set_title("Surge Protection Comparison — Peak Pressures", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3, axis="y")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _title_page(doc: Document, meta: dict) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ALLL WPS DESIGNER")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Pump Station Design Report")
    sr.font.size = Pt(16)
    sr.font.color.rgb = TEAL

    doc.add_paragraph()
    _hline(doc)
    doc.add_paragraph()

    stamp_rows = [
        ("Project",   meta.get("name",       "—")),
        ("Client",    meta.get("client",      "—")),
        ("Job No.",   meta.get("job_number",  "—")),
        ("Date",      meta.get("date",        str(_date.today()))),
        ("Engineer",  meta.get("engineer",    "—")),
        ("Status",    "Preliminary"),
    ]
    _kv_table(doc, stamp_rows, col_widths=(2.0, 4.0))

    if meta.get("notes"):
        doc.add_paragraph()
        _para(doc, "Design Notes:", bold=True)
        _para(doc, meta["notes"])

    _page_break(doc)


def _executive_summary(doc: Document, draft: dict) -> None:
    _heading(doc, "1. Executive Summary")

    meta  = draft.get("meta", {}) or {}
    hyd   = draft.get("hydraulicsResult") or {}
    pump  = draft.get("pumpResult")       or {}
    cw    = draft.get("clearwellResult")  or {}
    surge = draft.get("waterHammerResult") or {}
    us    = draft.get("unitSystem", "SI")
    Q     = draft.get("designFlow_m3h", 0) or 0

    _para(doc, (
        f"This report documents the engineering design calculations for the pump station "
        f"project '{meta.get('name', 'unnamed')}'. "
        f"The design flow is {_fmt(Q, 1)} m³/h and the system has been assessed for "
        f"hydraulic adequacy, pump selection, wet well capacity, and surge protection."
    ))

    doc.add_paragraph()
    rows: list[tuple[str, str]] = []

    if hyd:
        q_v, q_u = _us(us, Q, "flow")
        tdh_v, tdh_u = _us(us, hyd.get("tdh_m"), "head")
        rows += [
            ("Design Flow",          f"{_fmt(q_v, 1)} {q_u}"),
            ("Total Dynamic Head",   f"{_fmt(tdh_v)} {tdh_u}"),
            ("Pipe Velocity",        f"{_fmt(hyd.get('velocity_ms'), 3)} m/s"),
            ("Friction Head Loss",   f"{_fmt(hyd.get('friction_head_m'))} m"),
        ]

    ops = pump.get("operating_points", []) if pump else []
    if ops:
        op = ops[0]
        rows += [
            ("Pump Duty Point Q",    f"{_fmt(op.get('Q_m3h'))} m³/h"),
            ("Pump Duty Point H",    f"{_fmt(op.get('H_m'))} m"),
            ("Pump Efficiency",      f"{_fmt(op.get('eta_pct'), 1)} %"),
            ("Pump Shaft Power",     f"{_fmt(op.get('power_kW'))} kW"),
        ]
        npsh_ok = (op.get("npsh_margin_m") or 0) >= 0
        rows.append(("NPSH Status", "OK" if npsh_ok else "⚠ Check Required"))

    if cw and cw.get("active"):
        rows += [
            ("Operating Volume",     f"{_fmt(cw.get('operating_volume_m3'))} m³"),
            ("Detention Time",       f"{_fmt(cw.get('detention_time_min'), 1)} min"),
        ]

    if surge:
        cav = surge.get("cavitation_risk", False)
        rows += [
            ("Joukowsky ΔH",         f"{_fmt(surge.get('delta_H_joukowsky_m'))} m"),
            ("Surge Max Pressure",   f"{_fmt(surge.get('max_pressure_head_m'))} m head"),
            ("Cavitation Risk",      "YES — protection required" if cav else "No"),
        ]

    if rows:
        _kv_table(doc, rows)

    doc.add_paragraph()


def _basis_of_design(doc: Document, draft: dict) -> None:
    _heading(doc, "2. Basis of Design")
    _heading(doc, "2.1 Design Parameters", level=2)

    us   = draft.get("unitSystem", "SI")
    Q    = draft.get("designFlow_m3h", 0) or 0
    up   = draft.get("upstreamNode",   {}) or {}
    dn   = draft.get("downstreamNode", {}) or {}

    q_v,  q_u  = _us(us, Q,                                   "flow")
    hs_v, hs_u = _us(us, dn.get("elevation_m", 0) - up.get("elevation_m", 0), "head")

    rows = [
        ("Design flow Q",           f"{_fmt(q_v, 1)} {q_u}"),
        ("Suction node elevation",  f"{_fmt(up.get('elevation_m'), 2)} m"),
        ("Discharge node elevation",f"{_fmt(dn.get('elevation_m'), 2)} m"),
        ("Static head",             f"{_fmt(hs_v)} {hs_u}"),
        ("Unit system",             us),
    ]
    _kv_table(doc, rows)
    doc.add_paragraph()

    _heading(doc, "2.2 Pipeline Configuration", level=2)

    for side in ("suction", "discharge"):
        assy = draft.get(side, {}) or {}
        segs = assy.get("segments", []) or []
        k_sum = assy.get("accessories_K_sum", 0.0) or 0.0

        _para(doc, f"{side.capitalize()} assembly — {len(segs)} segment(s), ΣK = {_fmt(k_sum, 3)}", bold=True)

        if segs:
            headers  = ["#", "Material", "Diameter (mm)", "Length (m)"]
            data_rows = [
                [str(i + 1),
                 s.get("material", "—").upper(),
                 _fmt(s.get("diameter_mm"), 0),
                 _fmt(s.get("length_m"), 1)]
                for i, s in enumerate(segs)
            ]
            _col_table(doc, headers, data_rows, col_widths=[0.3, 1.8, 1.5, 1.5])

        doc.add_paragraph()


def _hydraulics_section(doc: Document, draft: dict) -> None:
    _heading(doc, "3. Hydraulic Analysis")

    us  = draft.get("unitSystem", "SI")
    Q   = draft.get("designFlow_m3h", 0) or 0
    hyd = draft.get("hydraulicsResult") or {}

    _heading(doc, "3.1 Friction Loss Method", level=2)
    _para(doc, (
        "Friction head losses are computed using the Darcy-Weisbach equation with "
        "the Colebrook-White implicit friction factor, iterated to convergence. "
        "Minor losses are calculated as ΣK × V²/(2g). "
        "Hazen-Williams coefficients are used where specified per segment."
    ))
    doc.add_paragraph()

    _heading(doc, "3.2 Key Equations", level=2)
    _para(doc, "Darcy-Weisbach friction head loss:", bold=True)
    _para(doc, "    hf = f · (L/D) · V²/(2g)", italic=True)
    _para(doc, "Colebrook-White (implicit):", bold=True)
    _para(doc, "    1/√f = −2 log₁₀(ε/(3.7D) + 2.51/(Re·√f))", italic=True)
    _para(doc, "Total Dynamic Head:", bold=True)
    _para(doc, "    TDH = Hstatic + Hf + Hm + ΔHpressure + ΔHvelocity", italic=True)
    doc.add_paragraph()

    if not hyd:
        _para(doc, "Hydraulic compute not yet run — results not available.", italic=True)
        return

    _heading(doc, "3.3 Head Budget", level=2)

    q_v, q_u   = _us(us, Q,                        "flow")
    vel_v, vel_u = _us(us, hyd.get("velocity_ms"),   "velocity")
    hs_v, hs_u  = _us(us, hyd.get("static_head_m"), "head")
    hf_v, hf_u  = _us(us, hyd.get("friction_head_m"), "head")
    hm_v, hm_u  = _us(us, hyd.get("minor_head_m"),  "head")
    tdh_v, tdh_u = _us(us, hyd.get("tdh_m"),        "head")

    rows = [
        ("Design flow Q",           f"{_fmt(q_v, 1)} {q_u}"),
        ("Pipe velocity V",         f"{_fmt(vel_v, 3)} {vel_u}"),
        ("Reynolds number Re",      f"{hyd.get('reynolds_number', 0):,.0f}"),
        ("Darcy friction factor f", f"{_fmt(hyd.get('friction_factor'), 5)}"),
        ("Static head Hs",          f"{_fmt(hs_v)} {hs_u}"),
        ("Friction head Hf",        f"{_fmt(hf_v)} {hf_u}"),
        ("Minor head Hm",           f"{_fmt(hm_v)} {hm_u}"),
        ("Total Dynamic Head TDH",  f"{_fmt(tdh_v)} {tdh_u}  ← Design duty"),
    ]
    _kv_table(doc, rows)
    doc.add_paragraph()

    sys_pts = hyd.get("system_curve", []) or []
    if sys_pts:
        _heading(doc, "3.4 System Curve Data Points", level=2)
        headers = ["Q (m³/h)", "H (m)"]
        data_rows = [[_fmt(p.get("Q_m3h"), 1), _fmt(p.get("H_m"))] for p in sys_pts]
        _col_table(doc, headers, data_rows, col_widths=[1.5, 1.5])
        doc.add_paragraph()

    fig_bytes = _fig_system_pump_curve(draft)
    if fig_bytes:
        _embed_figure(doc, fig_bytes,
                      "Figure 1 — Pump H-Q and system curve overlay with duty operating point(s)")
    doc.add_paragraph()


def _pump_section(doc: Document, draft: dict) -> None:
    _heading(doc, "4. Pump Analysis")

    pump = draft.get("pumpResult") or {}
    us   = draft.get("unitSystem", "SI")

    if not pump or not pump.get("active"):
        _para(doc, "Pump analysis not activated.", italic=True)
        return

    _heading(doc, "4.1 Operating Points", level=2)

    ops = pump.get("operating_points", []) or []
    if ops:
        headers = ["Pumps", "Q (m³/h)", "H (m)", "η (%)", "P (kW)", "NPSHr (m)", "NPSHa (m)", "Margin (m)", "Warnings"]
        data_rows = []
        for op in ops:
            q_v, _  = _us(us, op.get("Q_m3h"),  "flow")
            h_v, _  = _us(us, op.get("H_m"),    "head")
            data_rows.append([
                str(op.get("n_pumps", 1)),
                _fmt(q_v),
                _fmt(h_v),
                _fmt(op.get("eta_pct"), 1),
                _fmt(op.get("power_kW")),
                _fmt(op.get("npshr_m")),
                _fmt(op.get("npsha_m")),
                _fmt(op.get("npsh_margin_m")),
                "; ".join(op.get("warnings", [])) or "None",
            ])
        _col_table(doc, headers, data_rows,
                   col_widths=[0.5, 0.8, 0.7, 0.6, 0.7, 0.8, 0.8, 0.8, 1.5])
        doc.add_paragraph()

    _heading(doc, "4.2 Performance Curves", level=2)

    hq_pts  = pump.get("hq_curve",    []) or []
    eta_pts = pump.get("eta_curve",   []) or []
    p_pts   = pump.get("p_curve",     []) or []
    npshr   = pump.get("npshr_curve", []) or []

    if hq_pts:
        headers = ["Q (m³/h)", "H (m)", "η (%)", "P (kW)", "NPSHr (m)"]
        n = len(hq_pts)
        data_rows = []
        for i, hp in enumerate(hq_pts):
            eta_val = eta_pts[i].get("value") if i < len(eta_pts) else None
            p_val   = p_pts[i].get("value")   if i < len(p_pts)   else None
            nr_val  = npshr[i].get("value")   if i < len(npshr)   else None
            data_rows.append([
                _fmt(hp.get("Q_m3h"), 1),
                _fmt(hp.get("value")),
                _fmt(eta_val, 1),
                _fmt(p_val),
                _fmt(nr_val),
            ])
        _col_table(doc, headers, data_rows,
                   col_widths=[1.0, 1.0, 1.0, 1.0, 1.0])
        doc.add_paragraph()

    sp_curves = pump.get("speed_curves", []) or []
    if sp_curves:
        _heading(doc, "4.3 Variable Frequency Drive (VFD) — Affinity Law Curves", level=2)
        _para(doc, (
            "Affinity laws govern pump performance at reduced speed: "
            "Q ∝ n, H ∝ n², P ∝ n³. Speed-stepped H-Q curves are shown below."
        ))
        doc.add_paragraph()

    fig_eff = _fig_efficiency_curves(draft)
    if fig_eff:
        _embed_figure(doc, fig_eff,
                      "Figure 2 — Pump efficiency, shaft power, and NPSHr vs flow")
    doc.add_paragraph()

    warnings = pump.get("warnings", []) or []
    if pump.get("non_physical_fit"):
        warnings = ["Non-physical polynomial fit detected — curve data may be insufficient."] + warnings
    if warnings:
        _para(doc, "Pump Analysis Warnings:", bold=True)
        for w in warnings:
            _para(doc, f"  • {w}")
        doc.add_paragraph()


def _wetwell_section(doc: Document, draft: dict) -> None:
    _heading(doc, "5. Wet Well Sizing and Pump Cycling")

    cw = draft.get("clearwellResult") or {}
    us = draft.get("unitSystem", "SI")

    if not cw or not cw.get("active"):
        _para(doc, "Wet well analysis not activated.", italic=True)
        return

    _heading(doc, "5.1 Key Results", level=2)

    op_v, op_u = _us(us, cw.get("operating_volume_m3"), "volume")
    rows = [
        ("Operating volume (LWL → HWL)",  f"{_fmt(op_v)} {op_u}"),
        ("Detention time",                 f"{_fmt(cw.get('detention_time_min'), 1)} min"),
        ("Required detention",             f"{_fmt(cw.get('required_detention_min'), 1)} min"),
        ("Detention OK",                   "Yes" if cw.get("detention_ok") else "No / Not checked"),
    ]
    _kv_table(doc, rows)
    doc.add_paragraph()

    cr_list = cw.get("cycle_results", []) or []
    if cr_list:
        _heading(doc, "5.2 Pump Cycle Analysis (AWWA M32)", level=2)
        headers = ["Stage", "Label", "Q pump (m³/h)", "Q in (m³/h)", "t drain (s)", "t cycle (s)", "Cycles/hr", "OK"]
        data_rows = []
        for cr in cr_list:
            data_rows.append([
                str(cr.get("stage", "")),
                cr.get("label", ""),
                _fmt(cr.get("Q_pump_m3h"), 1),
                _fmt(cr.get("Q_in_m3h"),  1),
                _fmt(cr.get("t_drain_s"), 0),
                _fmt(cr.get("t_cycle_s"), 0),
                _fmt(cr.get("cycles_per_hour"), 2),
                "Yes" if cr.get("cycles_ok") else "No",
            ])
        _col_table(doc, headers, data_rows,
                   col_widths=[0.5, 1.0, 1.0, 1.0, 0.9, 0.9, 0.8, 0.5])
        doc.add_paragraph()

    vol_curve = cw.get("volume_curve", []) or []
    if vol_curve:
        _heading(doc, "5.3 Volume Curve", level=2)
        headers   = ["Level (m)", "Depth (m)", "Volume (m³)"]
        data_rows = [
            [_fmt(p.get("level_m")), _fmt(p.get("depth_m")), _fmt(p.get("volume_m3"))]
            for p in vol_curve
        ]
        _col_table(doc, headers, data_rows, col_widths=[1.5, 1.5, 1.5])
        doc.add_paragraph()

    warnings = cw.get("warnings", []) or []
    if warnings:
        _para(doc, "Wet Well Warnings:", bold=True)
        for w in warnings:
            _para(doc, f"  • {w}")
        doc.add_paragraph()


def _engineering_checks_section(doc: Document, draft: dict) -> None:
    _heading(doc, "6. Engineering Checks")

    hyd  = draft.get("hydraulicsResult") or {}
    pump = draft.get("pumpResult")       or {}
    cw   = draft.get("clearwellResult")  or {}
    us   = draft.get("unitSystem", "SI")

    checks: list[tuple[str, str, str]] = []  # (parameter, value, status)

    if hyd:
        vel  = hyd.get("velocity_ms", 0) or 0
        vel_v, vel_u = _us(us, vel, "velocity")
        v_ok = 0.5 <= vel <= 3.0
        checks.append(("Pipe velocity", f"{_fmt(vel_v, 3)} {vel_u}", "OK" if v_ok else "WARNING"))

        re   = hyd.get("reynolds_number", 0) or 0
        checks.append(("Reynolds number", f"{re:,.0f}", "OK" if re > 4000 else "INFO"))

    ops = pump.get("operating_points", []) if pump else []
    for op in ops:
        margin = op.get("npsh_margin_m")
        if margin is not None:
            s = "OK" if margin >= 0.6 else ("WARNING" if margin >= 0 else "CRITICAL")
            checks.append((f"NPSH margin (n={op.get('n_pumps',1)})",
                            f"{_fmt(margin)} m", s))
        eta = op.get("eta_pct")
        if eta is not None:
            checks.append((f"Pump efficiency (n={op.get('n_pumps',1)})",
                            f"{_fmt(eta, 1)} %", "OK" if eta >= 60 else "WARNING"))

    if cw and cw.get("active"):
        for cr in (cw.get("cycle_results", []) or []):
            s = "OK" if cr.get("cycles_ok") else "WARNING"
            checks.append((f"Pump cycles — {cr.get('label', '')}",
                            f"{_fmt(cr.get('cycles_per_hour'), 2)} /hr", s))
        dt_ok = cw.get("detention_ok")
        if dt_ok is not None:
            checks.append(("Detention time", f"{_fmt(cw.get('detention_time_min'), 1)} min",
                            "OK" if dt_ok else "WARNING"))

    if not checks:
        _para(doc, "No checks available — run analyses first.", italic=True)
        return

    COLOR_MAP = {"OK": "27AE60", "WARNING": "F39C12", "CRITICAL": "E74C3C", "INFO": "2980B9"}
    tbl = doc.add_table(rows=1 + len(checks), cols=3)
    tbl.style = "Table Grid"
    for j, h in enumerate(["Parameter", "Value", "Status"]):
        c = tbl.rows[0].cells[j]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        _set_cell_bg(c, "1F3864")
        _set_cell_borders(c)

    for i, (param, val, status) in enumerate(checks):
        row = tbl.rows[i + 1]
        row.cells[0].text = param
        row.cells[1].text = val
        row.cells[2].text = status
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_borders(cell)
            _set_cell_bg(cell, "F2F5F8" if i % 2 == 1 else "FFFFFF")
        status_cell = row.cells[2]
        _set_cell_bg(status_cell, COLOR_MAP.get(status, "FFFFFF"))
        if status != "OK":
            run = status_cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = WHITE

    doc.add_paragraph()


def _surge_section(doc: Document, draft: dict) -> None:
    _heading(doc, "7. Surge / Water Hammer Analysis")

    surge = draft.get("waterHammerResult") or {}
    moc   = draft.get("mocResult")         or {}
    wi    = draft.get("whatIfResult")      or {}
    us    = draft.get("unitSystem", "SI")

    _heading(doc, "7.1 Joukowsky (Mode A) — Instantaneous Analysis", level=2)
    _para(doc, (
        "The Joukowsky equation gives the maximum pressure transient for an "
        "instantaneous valve closure or pump trip: ΔH = a·ΔV / g. "
        "This is a conservative upper bound."
    ))
    doc.add_paragraph()

    if surge:
        rows = [
            ("Wave speed a",            f"{_fmt(surge.get('wave_speed_ms'), 0)} m/s"),
            ("Initial velocity V₀",     f"{_fmt(surge.get('V0_ms'), 3)} m/s"),
            ("Joukowsky ΔH",            f"{_fmt(surge.get('delta_H_joukowsky_m'))} m"),
            ("Joukowsky ΔP",            f"{_fmt(surge.get('delta_P_joukowsky_kPa'), 0)} kPa"),
            ("Operating head",          f"{_fmt(surge.get('H_operating_m'))} m"),
            ("Max transient head",      f"{_fmt(surge.get('max_pressure_head_m'))} m"),
            ("Min transient head",      f"{_fmt(surge.get('min_pressure_head_m'))} m"),
            ("Cavitation risk",         "YES" if surge.get("cavitation_risk") else "No"),
            ("Vacuum risk",             "YES" if surge.get("vacuum_risk")     else "No"),
            ("Characteristic time Tc",  f"{_fmt(surge.get('T_char_s'), 3)} s"),
        ]
        _kv_table(doc, rows)
        doc.add_paragraph()
    else:
        _para(doc, "Surge quick analysis not run.", italic=True)

    if moc:
        _heading(doc, "7.2 Method of Characteristics (MOC) — Full Transient", level=2)
        _para(doc, (
            "The Method of Characteristics solves the 1D transient pipe-flow PDEs on "
            "a characteristic grid. It captures wave reflections, attenuation, and "
            "node-by-node pressure histories."
        ))
        doc.add_paragraph()

        rows = [
            ("Pipe segments N",         str(moc.get("N", "—"))),
            ("Spatial step Δx",         f"{_fmt(moc.get('dx_m'))} m"),
            ("Time step Δt",            f"{_fmt(moc.get('dt_s'), 4)} s"),
            ("Courant number",          f"{_fmt(moc.get('courant'), 3)}"),
            ("Simulation duration",     f"{_fmt(moc.get('t_total_s'), 1)} s"),
            ("Global max head H_max",   f"{_fmt(moc.get('global_max_H_m'))} m"),
            ("Global min head H_min",   f"{_fmt(moc.get('global_min_H_m'))} m"),
            ("Cavitation locations",    str(len(moc.get("cavitation_x_m", []) or []))),
        ]
        _kv_table(doc, rows)
        doc.add_paragraph()

        env_fig = _fig_surge_envelope(draft)
        if env_fig:
            _embed_figure(doc, env_fig,
                          "Figure 3 — MOC pressure envelope: H_max and H_min vs distance")

        hist_fig = _fig_moc_time_history(draft)
        if hist_fig:
            _embed_figure(doc, hist_fig,
                          "Figure 4 — MOC head time histories at key observation nodes")

        doc.add_paragraph()

        obs = moc.get("observations", []) or []
        if obs:
            _heading(doc, "7.3 Time History Summary at Key Nodes", level=2)
            headers = ["Node", "x (m)", "H max (m)", "H min (m)"]
            data_rows = []
            for ob in obs:
                hist = ob.get("history", []) or []
                h_vals = [h.get("H_m", 0) for h in hist]
                data_rows.append([
                    ob.get("label", ""),
                    _fmt(ob.get("x_m")),
                    _fmt(max(h_vals) if h_vals else None),
                    _fmt(min(h_vals) if h_vals else None),
                ])
            _col_table(doc, headers, data_rows, col_widths=[2.0, 1.0, 1.2, 1.2])
            doc.add_paragraph()

    if wi and (wi.get("baseline") or wi.get("device_runs")):
        _heading(doc, "7.4 Surge Protection Device Comparison", level=2)

        baseline = wi.get("baseline") or {}
        runs     = wi.get("device_runs", []) or []
        headers  = ["Scenario", "H max (m)", "H min (m)", "ΔH reduction (m)", "Cavitation"]
        data_rows = []
        for r in [baseline] + runs:
            data_rows.append([
                r.get("label", "Baseline"),
                _fmt(r.get("global_max_H_m")),
                _fmt(r.get("global_min_H_m")),
                _fmt(r.get("max_surge_reduction_m")),
                "Yes" if r.get("cavitation_risk") else "No",
            ])
        _col_table(doc, headers, data_rows, col_widths=[2.0, 1.0, 1.0, 1.2, 1.0])
        doc.add_paragraph()

        comp_fig = _fig_protection_comparison(draft)
        if comp_fig:
            _embed_figure(doc, comp_fig,
                          "Figure 5 — Peak surge pressures: baseline vs. protection devices")

        notes = wi.get("assumption_notes", []) or []
        if notes:
            _para(doc, "Assumptions:", bold=True)
            for n in notes:
                _para(doc, f"  • {n}")
        doc.add_paragraph()


def _appendix_inputs(doc: Document, draft: dict) -> None:
    _page_break(doc)
    _heading(doc, "Appendix A — Complete Input Parameters")

    meta = draft.get("meta", {}) or {}
    rows = [
        ("Project name",   meta.get("name",      "—")),
        ("Client",         meta.get("client",     "—")),
        ("Job number",     meta.get("job_number", "—")),
        ("Date",           meta.get("date",       "—")),
        ("Engineer",       meta.get("engineer",   "—")),
        ("Unit system",    draft.get("unitSystem", "SI")),
        ("Design flow",    f"{_fmt(draft.get('designFlow_m3h'), 1)} m³/h"),
        ("Upstream elev.", f"{_fmt((draft.get('upstreamNode') or {}).get('elevation_m'), 2)} m"),
        ("Downstream elev.", f"{_fmt((draft.get('downstreamNode') or {}).get('elevation_m'), 2)} m"),
    ]
    _kv_table(doc, rows)
    doc.add_paragraph()

    for side in ("suction", "discharge"):
        assy = draft.get(side, {}) or {}
        segs = assy.get("segments", []) or []
        _para(doc, f"{side.capitalize()} segments:", bold=True)
        if segs:
            headers  = ["#", "Material", "Ø (mm)", "L (m)"]
            data_rows = [
                [str(i + 1), s.get("material", ""), _fmt(s.get("diameter_mm"), 0), _fmt(s.get("length_m"), 1)]
                for i, s in enumerate(segs)
            ]
            _col_table(doc, headers, data_rows, col_widths=[0.3, 2.0, 1.0, 1.0])
        doc.add_paragraph()


def _appendix_kvalue_ref(doc: Document) -> None:
    _heading(doc, "Appendix B — Pipe Roughness Reference", level=1)

    headers = ["Material", "Roughness ε (m)", "Typical HW-C"]
    data_rows = [
        ["PVC / HDPE / uPVC", "0.0000015",  "150"],
        ["Ductile iron (DICL)", "0.000250", "130"],
        ["Steel (unlined)",    "0.000046",   "120"],
        ["Asbestos cement",    "0.000030",   "140"],
        ["Concrete",           "0.000300",   "110"],
        ["GRP / FRP",          "0.000030",   "150"],
    ]
    _col_table(doc, headers, data_rows, col_widths=[2.5, 1.8, 1.5])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_document(draft: dict) -> Document:
    """
    Build a full engineering Word report from a ProjectDraft dict.

    Parameters
    ----------
    draft : dict
        Serialised ProjectDraft — the same payload used by the Excel exporter.

    Returns
    -------
    docx.Document
        Fully populated python-docx Document object ready for serialisation.
    """
    draft = draft or {}
    meta  = draft.get("meta", {}) or {}

    doc = Document()

    # Page margins (1 inch all round)
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title page
    _title_page(doc, meta)

    # Body sections
    _executive_summary(doc, draft)
    _basis_of_design(doc, draft)
    _hydraulics_section(doc, draft)
    _pump_section(doc, draft)
    _wetwell_section(doc, draft)
    _engineering_checks_section(doc, draft)
    _surge_section(doc, draft)

    # Appendices
    _appendix_inputs(doc, draft)
    _appendix_kvalue_ref(doc)

    return doc


def _doc_to_bytes(doc: Document) -> bytes:
    """Serialise a Document to a .docx byte string."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
